from __future__ import annotations

import base64
import binascii
import csv
import io
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from .chat_service import TextAttachment
from .config import settings
from .image_service import ImageAttachment


PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
MAX_OFFICE_ARCHIVE_FILES = 5_000
MAX_OFFICE_EXPANDED_BYTES = 512 * 1024 * 1024
VIDEO_MIME_BY_SUFFIX = {
    ".mp4": "video/mp4",
    ".mkv": "video/x-matroska",
    ".mov": "video/quicktime",
    ".webm": "video/webm",
    ".avi": "video/x-msvideo",
}
IMAGE_MIME_BY_SUFFIX = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


@dataclass(frozen=True)
class ParsedDocument:
    content: bytes
    mime_type: str
    text_attachment: TextAttachment
    images: list[ImageAttachment]


def decode_attachment_data_url(data_url: str, name: str) -> bytes:
    if not data_url.startswith("data:") or "," not in data_url:
        raise ValueError(f"附件数据无效：{name}")
    header, encoded = data_url.split(",", 1)
    if ";base64" not in header.lower():
        raise ValueError(f"附件不是有效的 base64 数据：{name}")
    try:
        content = base64.b64decode(encoded, validate=True)
    except (ValueError, binascii.Error) as exc:
        raise ValueError(f"附件数据损坏：{name}") from exc
    if not content:
        raise ValueError(f"附件内容为空：{name}")
    if len(content) > settings.agent_document_attachment_max_bytes:
        max_mb = settings.agent_document_attachment_max_bytes // (1024 * 1024)
        raise ValueError(f"附件超过 {max_mb}MB：{name}")
    return content


def _image_attachment(content: bytes, mime_type: str, source: str) -> ImageAttachment:
    encoded = base64.b64encode(content).decode("ascii")
    return ImageAttachment(
        data_url=f"data:{mime_type};base64,{encoded}",
        mime_type=mime_type,
        source=source,
        content=content,
    )


def _render_pdf_page(document, page_index: int, name: str) -> ImageAttachment:
    page = document[page_index]
    try:
        bitmap = page.render(scale=1.6)
        try:
            image = bitmap.to_pil().convert("RGB")
        finally:
            bitmap.close()
    finally:
        page.close()
    output = io.BytesIO()
    image.save(output, format="JPEG", quality=82, optimize=True)
    return _image_attachment(output.getvalue(), "image/jpeg", f"{name} 第 {page_index + 1} 页")


def _validate_office_archive(content: bytes, name: str) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_OFFICE_ARCHIVE_FILES:
                raise ValueError(f"Office 文档内部条目数超过 {MAX_OFFICE_ARCHIVE_FILES}：{name}")
            names: set[str] = set()
            expanded_size = 0
            for info in infos:
                normalized = info.filename.casefold()
                if normalized in names:
                    raise ValueError(f"Office 文档包含重复内部条目：{name}")
                if info.flag_bits & 0x1:
                    raise ValueError(f"Office 文档包含加密内部条目：{name}")
                names.add(normalized)
                expanded_size += info.file_size
                if expanded_size > MAX_OFFICE_EXPANDED_BYTES:
                    raise ValueError(f"Office 文档展开后超过 512MB：{name}")
    except zipfile.BadZipFile as exc:
        raise ValueError(f"不是有效的 Office 文档：{name}") from exc


def _parse_pdf(content: bytes, name: str) -> ParsedDocument:
    if not content.startswith(b"%PDF"):
        raise ValueError(f"不是有效的 PDF 文件：{name}")
    try:
        document = pdfium.PdfDocument(content)
    except Exception as exc:
        raise ValueError(f"无法打开 PDF，文件可能损坏或已加密：{name}") from exc

    try:
        page_count = len(document)
        if page_count > settings.agent_pdf_max_pages:
            raise ValueError(f"PDF 页数超过 {settings.agent_pdf_max_pages} 页：{name}")

        text_sections: list[str] = []
        scan_pages: list[int] = []
        for index in range(page_count):
            page = document[index]
            try:
                text_page = page.get_textpage()
                try:
                    text = re.sub(r"\x00+", "", text_page.get_text_bounded()).strip()
                finally:
                    text_page.close()
            finally:
                page.close()
            if len(re.sub(r"\s+", "", text)) >= 20:
                text_sections.append(f"## 第 {index + 1} 页\n{text}")
            else:
                scan_pages.append(index)

        selected_scan_pages = scan_pages[: settings.agent_document_vision_max_pages]
        images = [_render_pdf_page(document, index, name) for index in selected_scan_pages]
        notes: list[str] = []
        if selected_scan_pages:
            page_labels = "、".join(str(index + 1) for index in selected_scan_pages)
            notes.append(f"扫描页 {page_labels} 已作为图片附在本轮消息中，请结合图片读取。")
        if len(scan_pages) > len(selected_scan_pages):
            notes.append(
                f"另有 {len(scan_pages) - len(selected_scan_pages)} 个无文字页面未发送，"
                f"单次最多识别 {settings.agent_document_vision_max_pages} 个扫描页。"
            )

        extracted = "\n\n".join(text_sections)
        if notes:
            extracted = (extracted + "\n\n" if extracted else "") + "[PDF 读取说明]\n" + "\n".join(notes)
        if not extracted and not images:
            raise ValueError(f"PDF 中没有可读取的文字或页面：{name}")
        return ParsedDocument(
            content=content,
            mime_type=PDF_MIME_TYPE,
            text_attachment=TextAttachment(name=name, content=extracted, mime_type="text/plain"),
            images=images,
        )
    finally:
        document.close()


def _docx_images(content: bytes, name: str) -> list[ImageAttachment]:
    images: list[ImageAttachment] = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            media_names = sorted(item for item in archive.namelist() if item.startswith("word/media/"))
            for media_name in media_names[: settings.agent_document_vision_max_pages]:
                suffix = Path(media_name).suffix.lower()
                mime_type = IMAGE_MIME_BY_SUFFIX.get(suffix)
                if not mime_type:
                    continue
                info = archive.getinfo(media_name)
                image_limit = min(
                    settings.agent_document_attachment_max_bytes,
                    settings.qq_image_max_bytes,
                )
                if info.file_size > image_limit:
                    continue
                image_content = archive.read(media_name)
                try:
                    Image.open(io.BytesIO(image_content)).verify()
                except Exception:
                    continue
                images.append(_image_attachment(image_content, mime_type, f"{name} 内嵌图片"))
    except zipfile.BadZipFile as exc:
        raise ValueError(f"不是有效的 DOCX 文件：{name}") from exc
    return images


def _parse_docx(content: bytes, name: str) -> ParsedDocument:
    _validate_office_archive(content, name)
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(f"无法打开 DOCX，文件可能损坏或受到密码保护：{name}") from exc

    sections: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"[表格 {table_index}]\n" + "\n".join(rows))

    images = _docx_images(content, name)
    extracted = "\n\n".join(sections).strip()
    if images:
        extracted = (extracted + "\n\n" if extracted else "") + "[DOCX 读取说明]\n文档内嵌图片已附在本轮消息中。"
    if not extracted and not images:
        raise ValueError(f"DOCX 中没有可读取的文字、表格或图片：{name}")
    return ParsedDocument(
        content=content,
        mime_type=DOCX_MIME_TYPE,
        text_attachment=TextAttachment(name=name, content=extracted, mime_type="text/plain"),
        images=images,
    )


def _decode_tabular_text(content: bytes, name: str) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别表格文本编码：{name}")


def _table_preview(rows: list[list[object]], *, max_rows: int = 60, max_columns: int = 30) -> str:
    lines: list[str] = []
    for row in rows[:max_rows]:
        cells = [re.sub(r"\s+", " ", str(value or "")).strip()[:160] for value in row[:max_columns]]
        lines.append(" | ".join(cells))
    if len(rows) > max_rows:
        lines.append(f"……另有 {len(rows) - max_rows} 行未展开")
    return "\n".join(lines)


def _parse_delimited_table(content: bytes, name: str, delimiter: str) -> ParsedDocument:
    text = _decode_tabular_text(content, name)
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    if not rows:
        raise ValueError(f"表格中没有可读取的数据：{name}")
    max_columns = max((len(row) for row in rows), default=0)
    summary = (
        f"[表格摘要]\n文件：{name}\n总行数：{len(rows)}\n最大列数：{max_columns}\n\n"
        + _table_preview(rows)
    )
    return ParsedDocument(
        content=content,
        mime_type="text/csv" if delimiter == "," else "text/tab-separated-values",
        text_attachment=TextAttachment(name=name, content=summary, mime_type="text/plain"),
        images=[],
    )


def _parse_xlsx(content: bytes, name: str) -> ParsedDocument:
    _validate_office_archive(content, name)
    try:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise ValueError(f"无法打开 XLSX，文件可能损坏或受到密码保护：{name}") from exc
    sections: list[str] = []
    try:
        for sheet in workbook.worksheets[:8]:
            rows = [list(row) for row in sheet.iter_rows(values_only=True, max_row=200, max_col=40)]
            while rows and not any(value not in {None, ""} for value in rows[-1]):
                rows.pop()
            sections.append(
                f"## 工作表：{sheet.title}\n读取行数：{len(rows)}\n\n"
                + (_table_preview(rows) if rows else "空工作表")
            )
        if len(workbook.worksheets) > 8:
            sections.append(f"另有 {len(workbook.worksheets) - 8} 个工作表未展开。")
    finally:
        workbook.close()
    return ParsedDocument(
        content=content,
        mime_type=XLSX_MIME_TYPE,
        text_attachment=TextAttachment(
            name=name,
            content="[Excel 工作簿摘要]\n" + "\n\n".join(sections),
            mime_type="text/plain",
        ),
        images=[],
    )


def _ffmpeg_path() -> Path:
    candidates = [
        settings.voice_training_dir / "cache" / "bin" / "ffmpeg.exe",
        settings.voice_training_dir / "ffmpeg.exe",
    ]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    raise ValueError("缺少本地 FFmpeg，暂时不能解析视频关键帧。")


def _video_duration_seconds(ffmpeg: Path, source: Path) -> float:
    result = subprocess.run(
        [str(ffmpeg), "-hide_banner", "-i", str(source)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=30,
        creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0,
    )
    match = re.search(r"Duration:\s*(\d+):(\d+):(\d+(?:\.\d+)?)", result.stderr or "")
    if not match:
        return 0.0
    return int(match.group(1)) * 3600 + int(match.group(2)) * 60 + float(match.group(3))


def _parse_video(content: bytes, name: str, suffix: str) -> ParsedDocument:
    ffmpeg = _ffmpeg_path()
    settings.data_dir.mkdir(parents=True, exist_ok=True)
    images: list[ImageAttachment] = []
    subtitle_text = ""
    with tempfile.TemporaryDirectory(prefix="mio-video-", dir=str(settings.data_dir)) as temp_name:
        temp_dir = Path(temp_name)
        source = temp_dir / f"source{suffix}"
        source.write_bytes(content)
        duration = _video_duration_seconds(ffmpeg, source)
        if duration <= 0:
            timestamps = [0.0, 15.0, 45.0]
        else:
            timestamps = sorted({max(0.0, duration * ratio) for ratio in (0.03, 0.25, 0.5, 0.75, 0.95)})
        flags = subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0
        for index, timestamp in enumerate(timestamps[:6], start=1):
            output = temp_dir / f"frame-{index:02d}.jpg"
            result = subprocess.run(
                [
                    str(ffmpeg), "-hide_banner", "-loglevel", "error", "-y",
                    "-ss", f"{timestamp:.3f}", "-i", str(source), "-frames:v", "1",
                    "-vf", "scale='min(960,iw)':-2", "-q:v", "3", str(output),
                ],
                capture_output=True,
                timeout=45,
                creationflags=flags,
            )
            if result.returncode == 0 and output.is_file() and output.stat().st_size > 0:
                images.append(
                    _image_attachment(output.read_bytes(), "image/jpeg", f"{name} {timestamp:.1f} 秒关键帧")
                )

        subtitle = temp_dir / "subtitle.srt"
        result = subprocess.run(
            [
                str(ffmpeg), "-hide_banner", "-loglevel", "error", "-y", "-i", str(source),
                "-map", "0:s:0", "-f", "srt", str(subtitle),
            ],
            capture_output=True,
            timeout=45,
            creationflags=flags,
        )
        if result.returncode == 0 and subtitle.is_file():
            subtitle_text = subtitle.read_text(encoding="utf-8", errors="replace")[:40000].strip()

    if not images and not subtitle_text:
        raise ValueError(f"视频中没有提取到可读取的关键帧或字幕：{name}")
    duration_text = f"{duration:.1f} 秒" if duration > 0 else "未识别"
    text = (
        f"[视频解析摘要]\n文件：{name}\n时长：{duration_text}\n"
        f"关键帧：{len(images)} 张（临时提取，解析后已删除磁盘帧）"
    )
    if subtitle_text:
        text += "\n\n[内嵌字幕]\n" + subtitle_text
    else:
        text += "\n\n没有检测到可提取的内嵌字幕。"
    return ParsedDocument(
        content=content,
        mime_type=VIDEO_MIME_BY_SUFFIX[suffix],
        text_attachment=TextAttachment(name=name, content=text, mime_type="text/plain"),
        images=images,
    )


def parse_document_data_url(data_url: str, name: str) -> ParsedDocument:
    content = decode_attachment_data_url(data_url, name)
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(content, name)
    if suffix == ".docx":
        return _parse_docx(content, name)
    if suffix == ".csv":
        return _parse_delimited_table(content, name, ",")
    if suffix == ".tsv":
        return _parse_delimited_table(content, name, "\t")
    if suffix == ".xlsx":
        return _parse_xlsx(content, name)
    if suffix in VIDEO_MIME_BY_SUFFIX:
        return _parse_video(content, name, suffix)
    if suffix == ".doc":
        raise ValueError(f"暂不支持旧版 DOC，请另存为 DOCX：{name}")
    if suffix == ".xls":
        raise ValueError(f"暂不支持旧版 XLS，请另存为 XLSX：{name}")
    raise ValueError(f"暂不支持读取这个文档：{name}")
