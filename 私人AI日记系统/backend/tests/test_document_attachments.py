from __future__ import annotations

import base64
import io
import subprocess
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw

from app.config import settings
from app import document_service
from app.document_service import DOCX_MIME_TYPE, PDF_MIME_TYPE, _ffmpeg_path, parse_document_data_url
from app.routes.agent import AgentAttachmentRequest, _prepare_attachments


def _data_url(mime_type: str, content: bytes) -> str:
    encoded = base64.b64encode(content).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("测试文档", level=1)
    document.add_paragraph("这是 DOCX 正文。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "澪"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _scanned_pdf_bytes() -> bytes:
    image = Image.new("RGB", (600, 800), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 100), "SCANNED PDF PAGE", fill="black")
    output = io.BytesIO()
    image.save(output, format="PDF")
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "任务"
    sheet.append(["事项", "状态", "耗时"])
    sheet.append(["上下文压缩", "完成", 2.5])
    sheet.append(["视频解析", "进行中", 1])
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


class DocumentAttachmentTests(unittest.TestCase):
    def test_csv_is_parsed_as_structured_table(self) -> None:
        content = "事项,状态\n上下文压缩,完成\n视频解析,进行中\n".encode("utf-8")
        parsed = parse_document_data_url(_data_url("text/csv", content), "计划.csv")

        self.assertIn("总行数：3", parsed.text_attachment.content)
        self.assertIn("上下文压缩 | 完成", parsed.text_attachment.content)

    def test_tsv_is_parsed_as_structured_table(self) -> None:
        content = "事项\t状态\n上下文压缩\t完成\n视频解析\t进行中\n".encode("utf-8")
        parsed = parse_document_data_url(_data_url("text/tab-separated-values", content), "计划.tsv")

        self.assertIn("总行数：3", parsed.text_attachment.content)
        self.assertIn("视频解析 | 进行中", parsed.text_attachment.content)

    def test_xlsx_lists_sheets_and_rows(self) -> None:
        parsed = parse_document_data_url(
            _data_url("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", _xlsx_bytes()),
            "计划.xlsx",
        )

        self.assertIn("工作表：任务", parsed.text_attachment.content)
        self.assertIn("视频解析 | 进行中 | 1", parsed.text_attachment.content)

    def test_video_extracts_ephemeral_keyframes(self) -> None:
        try:
            ffmpeg = _ffmpeg_path()
        except ValueError:
            self.skipTest("当前测试环境没有 FFmpeg")
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "sample.mp4"
            subprocess.run(
                [
                    str(ffmpeg), "-hide_banner", "-loglevel", "error", "-y",
                    "-f", "lavfi", "-i", "testsrc=size=320x180:rate=12:duration=2",
                    "-pix_fmt", "yuv420p", str(source),
                ],
                check=True,
                timeout=45,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0,
            )
            original_data_dir = settings.data_dir
            object.__setattr__(settings, "data_dir", Path(temp_dir))
            try:
                parsed = parse_document_data_url(
                    _data_url("video/mp4", source.read_bytes()),
                    "测试视频.mp4",
                )
                leftovers = list(Path(temp_dir).glob("mio-video-*"))
            finally:
                object.__setattr__(settings, "data_dir", original_data_dir)

        self.assertGreaterEqual(len(parsed.images), 1)
        self.assertIn("临时提取", parsed.text_attachment.content)
        self.assertEqual(leftovers, [])

    def test_video_extracts_embedded_subtitles(self) -> None:
        try:
            ffmpeg = _ffmpeg_path()
        except ValueError:
            self.skipTest("当前测试环境没有 FFmpeg")
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            subtitle = root / "subtitle.srt"
            subtitle.write_text(
                "1\n00:00:00,000 --> 00:00:01,500\n澪测试字幕\n",
                encoding="utf-8",
            )
            source = root / "subtitled.mkv"
            subprocess.run(
                [
                    str(ffmpeg), "-hide_banner", "-loglevel", "error", "-y",
                    "-f", "lavfi", "-i", "color=c=white:size=320x180:rate=12:duration=2",
                    "-i", str(subtitle), "-c:v", "mpeg4", "-c:s", "srt", "-shortest", str(source),
                ],
                check=True,
                timeout=45,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0,
            )
            original_data_dir = settings.data_dir
            object.__setattr__(settings, "data_dir", root)
            try:
                parsed = parse_document_data_url(
                    _data_url("video/x-matroska", source.read_bytes()),
                    "字幕视频.mkv",
                )
                leftovers = list(root.glob("mio-video-*"))
            finally:
                object.__setattr__(settings, "data_dir", original_data_dir)

        self.assertIn("[内嵌字幕]", parsed.text_attachment.content)
        self.assertIn("澪测试字幕", parsed.text_attachment.content)
        self.assertEqual(leftovers, [])

    def test_docx_extracts_paragraphs_and_tables(self) -> None:
        parsed = parse_document_data_url(_data_url(DOCX_MIME_TYPE, _docx_bytes()), "测试.docx")

        self.assertIn("这是 DOCX 正文", parsed.text_attachment.content)
        self.assertIn("项目 | 澪", parsed.text_attachment.content)
        self.assertEqual(parsed.mime_type, DOCX_MIME_TYPE)

    def test_office_document_rejects_excessive_internal_entries(self) -> None:
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", "one")
            archive.writestr("word/document.xml", "two")
        with patch.object(document_service, "MAX_OFFICE_ARCHIVE_FILES", 1):
            with self.assertRaisesRegex(ValueError, "内部条目数超过"):
                parse_document_data_url(_data_url(DOCX_MIME_TYPE, output.getvalue()), "异常.docx")

    def test_office_document_rejects_excessive_expanded_size(self) -> None:
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", "expanded")
        with patch.object(document_service, "MAX_OFFICE_EXPANDED_BYTES", 4):
            with self.assertRaisesRegex(ValueError, "展开后超过"):
                parse_document_data_url(_data_url(DOCX_MIME_TYPE, output.getvalue()), "压缩炸弹.docx")

    def test_scanned_pdf_is_rendered_for_vision(self) -> None:
        parsed = parse_document_data_url(_data_url(PDF_MIME_TYPE, _scanned_pdf_bytes()), "扫描.pdf")

        self.assertIn("扫描页 1", parsed.text_attachment.content)
        self.assertEqual(len(parsed.images), 1)
        self.assertEqual(parsed.images[0].mime_type, "image/jpeg")

    def test_document_attachment_archives_original_and_exposes_scan_image(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original = settings.agent_attachment_dir
            object.__setattr__(settings, "agent_attachment_dir", Path(temp_dir))
            try:
                images, text_files, metadata = _prepare_attachments([
                    AgentAttachmentRequest(
                        kind="document",
                        name="扫描.pdf",
                        mime_type=PDF_MIME_TYPE,
                        data_url=_data_url(PDF_MIME_TYPE, _scanned_pdf_bytes()),
                    )
                ])
            finally:
                object.__setattr__(settings, "agent_attachment_dir", original)

        self.assertEqual(len(images), 1)
        self.assertEqual(len(text_files), 1)
        self.assertEqual(metadata[0]["kind"], "document")
        self.assertEqual(metadata[0]["mime_type"], PDF_MIME_TYPE)

    def test_bmp_image_is_normalised_for_vision(self) -> None:
        output = io.BytesIO()
        Image.new("RGB", (24, 24), "navy").save(output, format="BMP")
        with tempfile.TemporaryDirectory() as temp_dir:
            original = settings.agent_attachment_dir
            object.__setattr__(settings, "agent_attachment_dir", Path(temp_dir))
            try:
                images, text_files, metadata = _prepare_attachments([
                    AgentAttachmentRequest(
                        kind="image",
                        name="测试图片.bmp",
                        mime_type="image/bmp",
                        data_url=_data_url("image/bmp", output.getvalue()),
                    )
                ])
            finally:
                object.__setattr__(settings, "agent_attachment_dir", original)

        self.assertEqual(text_files, [])
        self.assertEqual(len(images), 1)
        self.assertEqual(images[0].mime_type, "image/jpeg")
        self.assertEqual(metadata[0]["kind"], "image")

    def test_unknown_file_is_archived_with_unreadable_note(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original = settings.agent_attachment_dir
            object.__setattr__(settings, "agent_attachment_dir", Path(temp_dir))
            try:
                images, text_files, metadata = _prepare_attachments([
                    AgentAttachmentRequest(
                        kind="file",
                        name="资料.zip",
                        mime_type="application/zip",
                        data_url=_data_url("application/zip", b"not-a-real-zip"),
                    )
                ])
                archived = Path(temp_dir) / metadata[0]["url"].split("/")[-2] / Path(metadata[0]["url"]).name
                archived_exists = archived.exists()
            finally:
                object.__setattr__(settings, "agent_attachment_dir", original)

        self.assertEqual(images, [])
        self.assertEqual(len(text_files), 1)
        self.assertIn("没有对应的正文解析器", text_files[0].content)
        self.assertEqual(metadata[0]["kind"], "file")
        self.assertTrue(archived_exists)


if __name__ == "__main__":
    unittest.main()
